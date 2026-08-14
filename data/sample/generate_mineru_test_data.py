"""
MinerU 测试数据生成器
=====================
构造不规则含表格等多模态数据的测试合同 DOCX 文件.

测试数据覆盖以下多模态元素:
  1. 标题 (一级/二级标题)
  2. 正文段落 (合同条款文字)
  3. 表格 (付款计划表 / 交付物清单表)
  4. 印章/签名占位符 (文本标注)
  5. 混合格式 (中英文混排 / 编号列表)

运行方式:
  python data/sample/generate_mineru_test_data.py
  → 输出: data/sample/MinerU测试合同_多模态.docx
  → 输出: data/sample/MinerU测试合同_多模态.txt (纯文本对照版)
"""
# 📜 代码文字逻辑解析
# 本脚本用于生成 MinerU 文档解析的测试数据.
# 测试合同包含多模态元素(文字+表格+印章占位), 用于验证:
# 1) MinerU 能正确提取表格结构
# 2) 降级模式能正确处理纯文本
# 3) 后续规则层能从结构化 JSON 中正则提取金额/日期等字段

import os
import sys

# 尝试导入 python-docx, 不可用时给出提示
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx 未安装, 无法生成 DOCX 测试数据")
    print("   请运行: pip install python-docx")
    print("   将仅生成 TXT 纯文本对照版")


def generate_test_contract_docx(output_path: str):
    """
    生成包含多模态元素的测试合同 DOCX 文件.

    包含:
      - 标题 (一级/二级)
      - 正文段落 (甲乙方信息 / 合同标的 / 违约责任)
      - 付款计划表 (3列5行)
      - 交付物清单表 (4列4行)
      - 印章/签名占位符
    """
    if not HAS_DOCX:
        return False

    doc = Document()

    # ============== 设置默认字体 ==============
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # ============== 一级标题 ==============
    title = doc.add_heading('电脑设备采购合同', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============== 甲乙方信息 ==============
    doc.add_heading('一、合同当事人', level=2)

    p = doc.add_paragraph()
    p.add_run('甲方（采购方）：').bold = True
    p.add_run('北京智慧科技有限公司')

    p = doc.add_paragraph()
    p.add_run('统一社会信用代码：').bold = True
    p.add_run('91110108MA01ABC123')

    p = doc.add_paragraph()
    p.add_run('法定代表人：').bold = True
    p.add_run('张三')

    p = doc.add_paragraph()
    p.add_run('乙方（供货方）：').bold = True
    p.add_run('深圳创新科技股份有限公司')

    p = doc.add_paragraph()
    p.add_run('统一社会信用代码：').bold = True
    p.add_run('91440300MA5G2XY456')

    p = doc.add_paragraph()
    p.add_run('法定代表人：').bold = True
    p.add_run('李四')

    # ============== 合同标的 ==============
    doc.add_heading('二、合同标的', level=2)
    doc.add_paragraph(
        '甲方向乙方采购以下电脑设备，具体型号、数量、单价如下表所示。'
        '乙方应按照合同约定的质量标准、交货时间及交付地点向甲方交付货物。'
    )

    # ============== 交付物清单表 (多模态元素 1: 表格) ==============
    doc.add_heading('2.1 交付物清单', level=3)
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Light Grid Accent 1'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ['序号', '产品名称', '数量（台）', '单价（元）']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 数据行
    data = [
        ['1', '联想 ThinkPad X1 Carbon 笔记本电脑', '50', '12,800'],
        ['2', '戴尔 U2723QE 27英寸4K显示器', '50', '3,500'],
        ['3', '罗技 MX Master 3S 无线鼠标', '100', '599'],
        ['4', '合计总价', '—', '815,900'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table1.rows[row_idx].cells[col_idx].text = value

    # ============== 付款计划表 (多模态元素 2: 表格) ==============
    doc.add_heading('2.2 付款计划', level=3)
    doc.add_paragraph('甲方应按以下付款计划向乙方支付合同总价款：')

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Light Grid Accent 1'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ['付款阶段', '付款比例', '付款金额（元）', '付款条件']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    payment_data = [
        ['第一期（预付款）', '30%', '244,770', '合同签订后5个工作日内'],
        ['第二期（货到付款）', '40%', '326,360', '全部设备交付并验收合格后5个工作日内'],
        ['第三期（质保金）', '30%', '244,770', '质保期（12个月）届满后5个工作日内'],
        ['合计', '100%', '815,900', '—'],
    ]
    for row_idx, row_data in enumerate(payment_data, 1):
        for col_idx, value in enumerate(row_data):
            table2.rows[row_idx].cells[col_idx].text = value

    # ============== 违约责任 ==============
    doc.add_heading('三、违约责任', level=2)
    doc.add_paragraph(
        '3.1 乙方未按合同约定时间交货的，每逾期一日，应向甲方支付合同总价的千分之三（3‰）作为违约金，'
        '违约金累计不超过合同总价的百分之十（10%）。'
    )
    doc.add_paragraph(
        '3.2 甲方未按合同约定时间付款的，每逾期一日，应向乙方支付应付未付金额的千分之三（3‰）作为违约金。'
    )
    doc.add_paragraph(
        '3.3 乙方交付的设备不符合合同约定的质量标准的，甲方有权拒绝接收并要求乙方在10日内更换。'
        '乙方未在限期内更换的，甲方有权解除合同并要求乙方赔偿损失。'
    )

    # ============== 争议解决 ==============
    doc.add_heading('四、争议解决', level=2)
    doc.add_paragraph(
        '本合同履行过程中发生的争议，双方应首先通过友好协商解决；'
        '协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。'
    )

    # ============== 印章/签名占位符 (多模态元素 3) ==============
    doc.add_heading('五、签署信息', level=2)

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'

    sig_table.rows[0].cells[0].text = '甲方（盖章）：北京智慧科技有限公司\n\n[此处为甲方公章位置]\n\n法定代表人/授权代表（签字）：____________\n\n签署日期：2026年__月__日'
    sig_table.rows[0].cells[1].text = '乙方（盖章）：深圳创新科技股份有限公司\n\n[此处为乙方公章位置]\n\n法定代表人/授权代表（签字）：____________\n\n签署日期：2026年__月__日'

    # ============== 保存 ==============
    doc.save(output_path)
    print(f"✅ DOCX 测试合同已生成: {output_path}")
    return True


def generate_test_contract_txt(output_path: str):
    """
    生成纯文本对照版测试合同(无表格, 文字平铺).
    用于 MinerU 不可用时的降级测试.
    """
    content = """电脑设备采购合同

一、合同当事人
甲方（采购方）：北京智慧科技有限公司
统一社会信用代码：91110108MA01ABC123
法定代表人：张三
乙方（供货方）：深圳创新科技股份有限公司
统一社会信用代码：91440300MA5G2XY456
法定代表人：李四

二、合同标的
甲方向乙方采购以下电脑设备，具体型号、数量、单价如下表所示。
乙方应按照合同约定的质量标准、交货时间及交付地点向甲方交付货物。

2.1 交付物清单
序号\t产品名称\t数量（台）\t单价（元）
1\t联想 ThinkPad X1 Carbon 笔记本电脑\t50\t12,800
2\t戴尔 U2723QE 27英寸4K显示器\t50\t3,500
3\t罗技 MX Master 3S 无线鼠标\t100\t599
4\t合计总价\t—\t815,900

2.2 付款计划
甲方应按以下付款计划向乙方支付合同总价款：
付款阶段\t付款比例\t付款金额（元）\t付款条件
第一期（预付款）\t30%\t244,770\t合同签订后5个工作日内
第二期（货到付款）\t40%\t326,360\t全部设备交付并验收合格后5个工作日内
第三期（质保金）\t30%\t244,770\t质保期（12个月）届满后5个工作日内
合计\t100%\t815,900\t—

三、违约责任
3.1 乙方未按合同约定时间交货的，每逾期一日，应向甲方支付合同总价的千分之三（3‰）作为违约金，违约金累计不超过合同总价的百分之十（10%）。
3.2 甲方未按合同约定时间付款的，每逾期一日，应向乙方支付应付未付金额的千分之三（3‰）作为违约金。
3.3 乙方交付的设备不符合合同约定的质量标准的，甲方有权拒绝接收并要求乙方在10日内更换。乙方未在限期内更换的，甲方有权解除合同并要求乙方赔偿损失。

四、争议解决
本合同履行过程中发生的争议，双方应首先通过友好协商解决；协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。

五、签署信息
甲方（盖章）：北京智慧科技有限公司 [此处为甲方公章位置] 法定代表人/授权代表（签字）：____________ 签署日期：2026年__月__日
乙方（盖章）：深圳创新科技股份有限公司 [此处为乙方公章位置] 法定代表人/授权代表（签字）：____________ 签署日期：2026年__月__日
"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✅ TXT 测试合同已生成: {output_path}")


if __name__ == "__main__":
    # 输出目录
    output_dir = os.path.dirname(os.path.abspath(__file__))

    # 生成 DOCX (需要 python-docx)
    docx_path = os.path.join(output_dir, "MinerU测试合同_多模态.docx")
    if HAS_DOCX:
        generate_test_contract_docx(docx_path)
    else:
        print("⚠️ 跳过 DOCX 生成(python-docx 未安装)")

    # 生成 TXT (总是生成, 用于降级测试)
    txt_path = os.path.join(output_dir, "MinerU测试合同_多模态.txt")
    generate_test_contract_txt(txt_path)

    print(f"\n📋 测试数据生成完成, 可用于:")
    print(f"   1. MinerU 解析测试: 上传 DOCX 文件到 doc_extract_mineru_node")
    print(f"   2. 降级模式测试: 使用 TXT 文件验证纯文本解析")
    print(f"   3. 规则层测试: 从结构化 JSON 中正则提取金额/日期/百分比")
